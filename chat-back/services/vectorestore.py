import json
import docx
import boto3
from fnmatch import fnmatchcase
from typing import Optional, Dict, Any, List, Tuple
from io import BytesIO
from docx import Document as DocxDocument
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
# from .database import add_filename_to_metadata
from langchain_community.embeddings import HuggingFaceEmbeddings
from app.models import init_metadata_db
from docx.oxml.ns import qn
import sqlite3



current_user = "TEST2"


model_id = 'intfloat/multilingual-e5-large'
# model_id = 'intfloat/multilingual-e5-base'
model_kwargs = {'device': 'cpu'} # Настройка для использования CPU (можно переключить на GPU)
# model_kwargs = {'device': 'cuda'}
embeddings = HuggingFaceEmbeddings(
    model_name=model_id,
    model_kwargs=model_kwargs
)



session = boto3.session.Session()
s3_client = session.client(
    service_name='s3',
    endpoint_url='https://storage.yandexcloud.net',
    aws_access_key_id='YCAJEt7ilkMDiPuuZA--Sgb1H',
    aws_secret_access_key='YCOJE46MLMRlPll_kl6oIllqvT7P7S65E4QohXLZ',
)


init_metadata_db()
CHROMA_PATH = f'./chroma/{current_user}/primary/'


def add_filename_to_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''INSERT INTO uploaded_docs (global_source, filename) values ('{source}', '{filename}') ; ''')

class Document:
    def __init__(self, source: str, page_content: str, metadata: Optional[Dict[str, Any]] = None):
        self.source = source
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {'source': source}


def get_uploaded_filenames(source) -> List[str]:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT filename FROM uploaded_docs WHERE global_source = ?", (source,))
        rows = cursor.fetchall()
    filenames = [row[0] for row in rows]
    return filenames


# def load_s3_files(bucket: str, prefix: str, suffix: str) -> List[str]:
#     """List files in a given S3 bucket with a specified prefix and suffix."""
#     try:
#         response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)
#         files = [content['Key'] for content in response.get('Contents', []) if content['Key'].endswith(suffix)]
#         if not files:
#             print(f"No files found in bucket {bucket} with prefix {prefix} and suffix {suffix}")
#         else:
#             print(f"Files found in bucket {bucket} with prefix {prefix} and suffix {suffix}: {files}")
#         return files
#     except Exception as e:
#         print(f"Error listing files in bucket {bucket} with prefix {prefix} and suffix {suffix}: {e}")
#         return []
#
#
# def load_docx_new(source, bucket: str) -> List[Document]:
#     import os
#     prefix = 'TEST1/docx/'
#     suffix = '.docx'
#     files = load_s3_files(bucket, prefix, suffix)
#     uniq_files = get_uploaded_filenames(source) or []
#
#     docs = []
#     for file in files:
#         if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
#             try:
#                 obj = s3_client.get_object(Bucket=bucket, Key=file)
#                 content = obj['Body'].read()
#
#                 # Используем BytesIO для чтения содержимого файла как бинарного потока
#                 doc_stream = BytesIO(content)
#                 doc = DocxDocument(doc_stream)
#
#                 # Извлекаем текст из документа docx
#                 full_text = []
#                 image_counter = 1
#
#                 # Получаем имя файла без расширения и создаем соответствующую папку
#                 filename_without_extension = os.path.splitext(os.path.basename(file))[0]
#                 image_folder = filename_without_extension  # Используем оригинальное имя файла для папки
#
#                 for para in doc.paragraphs:
#                     # Обработка параграфов для создания ссылок на изображения
#                     para_text = para.text
#                     for run in para.runs:
#                         for drawing in run.element.findall('.//a:blip', namespaces={
#                             'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
#                             image_rId = drawing.get(qn('r:embed'))
#                             image_part = doc.part.related_parts[image_rId]
#                             image_filename = f'image_{image_counter:02d}.{image_part.content_type.split("/")[-1]}'
#                             image_counter += 1
#
#                             # Загрузка изображения в бакет Яндекса
#                             img_content = image_part.blob
#                             s3_image_key = f"TEST1/images/{image_folder}/{image_filename}"
#                             s3_client.put_object(
#                                 Bucket=bucket,
#                                 Key=s3_image_key,
#                                 Body=img_content,
#                                 ContentDisposition='inline',
#                                 ContentType=image_part.content_type
#                             )
#
#                             # Генерация URL для изображения
#                             s3_image_url = f"https://storage.yandexcloud.net/{bucket}/{s3_image_key}"
#                             para_text += f'\n{s3_image_url}'
#                     full_text.append(para_text)
#                 content = '\n'.join(full_text)
#
#                 docs.append(Document(source=file, page_content=content))
#             except Exception as e:
#                 print(f"Error reading docx file {file}: {e}")
#
#     return docs if docs else None
#
#
# def load_txts(source, bucket: str) -> List[Document]:
#     prefix = f'{current_user}/txt/'
#     suffix = '.txt'
#     files = load_s3_files(bucket, prefix, suffix)
#     uniq_files = get_uploaded_filenames(source) or []
#
#     docs = []
#     for file in files:
#         if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
#             try:
#                 obj = s3_client.get_object(Bucket=bucket, Key=file)
#                 content = obj['Body'].read().decode('utf-8')
#                 docs.append(Document(source=file, page_content=content))
#             except Exception as e:
#                 print(f"Error reading txt file {file}: {e}")
#
#     return docs if docs else None
#
#
# def load_jsons(source, bucket: str) -> Tuple[List[Document], List[dict]]:
#     prefix = f'{current_user}/json/'
#     suffix = '.json'
#     files = load_s3_files(bucket, prefix, suffix)
#     uniq_files = get_uploaded_filenames(source) or []
#
#     json_docs, json_metadata = [], []
#     for file in files:
#         if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
#             try:
#                 obj = s3_client.get_object(Bucket=bucket, Key=file)
#                 content = json.loads(obj['Body'].read().decode('utf-8'))
#                 json_docs.append(content)
#                 json_metadata.append({'source': file})
#             except Exception as e:
#                 print(f"Error reading json file {file}: {e}")
#
#     return (json_docs, json_metadata) if json_docs else (None, None)
#
#
#
# def load_csvs(source, bucket: str) -> List[Document]:
#     prefix = f'{current_user}/csv/'
#     suffix = '.csv'
#     files = load_s3_files(bucket, prefix, suffix)
#     uniq_files = get_uploaded_filenames(source) or []
#
#     docs = []
#     for file in files:
#         if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
#             try:
#                 obj = s3_client.get_object(Bucket=bucket, Key=file)
#                 content = obj['Body'].read().decode('utf-8')
#                 docs.append(Document(source=file, page_content=content))
#             except Exception as e:
#                 print(f"Error reading txt file {file}: {e}")
#
#     return docs if docs else None
#
#
# def load_documents(global_source, bucket: str, file_types: List[str]) -> dict:
#     """
#     Загружаем документы в зависимости от типа документа из Yandex S3
#     """
#     all_docs = {'txt': None, 'json': None, 'json_metadata': None, 'docx': None}
#     if 'txt' in file_types:
#         txt_docs = load_txts(global_source, bucket)
#         all_docs['txt'] = txt_docs
#     if 'json' in file_types:
#         json_docs, json_metadata = load_jsons(global_source, bucket)
#         all_docs['json'] = json_docs
#         all_docs['json_metadata'] = json_metadata
#     if 'docx' in file_types:
#         docx_docs = load_docx_new(global_source, bucket)
#         all_docs['docx'] = docx_docs
#     if 'csv' in file_types:
#         csv_docs = load_csvs(global_source, bucket)
#         all_docs['csv'] = csv_docs
#     return all_docs
#
#
# # Пример использования
# DATA_BUCKET = 'utlik'
# DOCS = load_documents('s3', DATA_BUCKET, ['txt', 'json', 'docx', 'csv'])

def list_local_files(directory: str, suffix: str) -> List[str]:
    import os
    """Собирает список файлов с определенным расширением в указанной директории."""
    files = []
    for root, _, filenames in os.walk(directory):
        for filename in filenames:
            if filename.endswith(suffix):
                files.append(os.path.join(root, filename))
    return files

def load_docx_local(directory: str) -> List[Document]:
    """Загружает файлы .docx из локальной директории."""
    files = list_local_files(directory, '.docx')
    docs = []
    for file in files:
        try:
            doc = DocxDocument(file)

            # Извлекаем текст из документа docx
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            content = '\n'.join(full_text)
            docs.append(Document(source=file, page_content=content))
        except Exception as e:
            print(f"Error reading docx file {file}: {e}")

    return docs

def load_txt_local(directory: str) -> List[Document]:
    """Загружает текстовые файлы (.txt) из локальной директории."""
    files = list_local_files(directory, '.txt')
    docs = []
    for file in files:
        try:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            docs.append(Document(source=file, page_content=content))
        except Exception as e:
            print(f"Error reading txt file {file}: {e}")

    return docs

def load_json_local(directory: str) -> Tuple[List[Document], List[dict]]:
    """Загружает JSON файлы из локальной директории."""
    files = list_local_files(directory, '.json')
    json_docs, json_metadata = [], []
    for file in files:
        try:
            with open(file, 'r', encoding='utf-8') as f:
                content = json.load(f)
            json_docs.append(content)
            json_metadata.append({'source': file})
        except Exception as e:
            print(f"Error reading json file {file}: {e}")

    return json_docs, json_metadata

def load_csv_local(directory: str) -> List[Document]:
    """Загружает CSV файлы из локальной директории."""
    files = list_local_files(directory, '.csv')
    docs = []
    for file in files:
        try:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            docs.append(Document(source=file, page_content=content))
        except Exception as e:
            print(f"Error reading csv file {file}: {e}")

    return docs


def load_documents_local(base_directory: str, file_types: List[str]) -> dict:
    """
    Загружаем документы из локальной папки.
    """
    all_docs = {'txt': None, 'json': None, 'json_metadata': None, 'docx': None, 'csv': None}
    if 'txt' in file_types:
        all_docs['txt'] = load_txt_local(base_directory)
    if 'json' in file_types:
        json_docs, json_metadata = load_json_local(base_directory)
        all_docs['json'] = json_docs
        all_docs['json_metadata'] = json_metadata
    if 'docx' in file_types:
        all_docs['docx'] = load_docx_local(base_directory)
    if 'csv' in file_types:
        all_docs['csv'] = load_csv_local(base_directory)
    return all_docs


# Пример использования
BASE_DIRECTORY = './upload_files'
FILE_TYPES = ['txt', 'json', 'docx', 'csv']

DOCS = load_documents_local(BASE_DIRECTORY, FILE_TYPES)


from typing import List

def split_docs_to_chunks(documents: dict, file_types: List[str], chunk_size=3000, chunk_overlap=100):
    all_chunks = []
    if 'txt' in file_types and documents['txt'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['txt']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    if 'json' in file_types and documents['json'] is not None:
        json_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        json_chunks = json_splitter.create_documents([json.dumps(doc, ensure_ascii=False) for doc in documents['json']],
                                                     metadatas=documents['json_metadata'])
        all_chunks.extend(json_chunks)

    if 'docx' in file_types and documents['docx'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['docx']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    if 'csv' in file_types and documents['csv'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['csv']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    return all_chunks


# def split_docs_to_chunks(documents: dict, file_types: List[str]):
#     all_chunks = []
#
#     def split_by_braces_and_comma(text):
#         # Разбиваем текст после каждого десятого сочетания '};'
#         pattern = r"(},)"
#         matches = list(re.finditer(pattern, text))
#
#         # Индексы, где нужно разрезать
#         split_indices = [match.end() for i, match in enumerate(matches) if (i + 1) % 8 == 0]
#
#         # Если разрезов не найдено, возвращаем оригинальный текст как один чанк
#         if not split_indices:
#             return [text]
#
#         # Создаем чанки по этим индексам
#         start_idx = 0
#         chunks = []
#         for idx in split_indices:
#             chunks.append(text[start_idx:idx].strip())
#             start_idx = idx
#
#         # Добавляем оставшуюся часть текста
#         chunks.append(text[start_idx:].strip())
#         return chunks
#
#     if 'txt' in file_types and documents['txt'] is not None:
#         for doc in documents['txt']:
#             chunks = split_by_braces_and_comma(doc.page_content)
#             for chunk in chunks:
#                 all_chunks.append(Document(source=doc.source, page_content=chunk, metadata=doc.metadata))
#
#     if 'json' in file_types and documents['json'] is not None:
#         for idx, doc in enumerate(documents['json']):
#             text = json.dumps(doc, ensure_ascii=False)
#             chunks = split_by_braces_and_comma(text)
#             for chunk in chunks:
#                 all_chunks.append(Document(source=documents['json_metadata'][idx]['source'], page_content=chunk))
#
#     if 'docx' in file_types and documents['docx'] is not None:
#         for doc in documents['docx']:
#             chunks = split_by_braces_and_comma(doc.page_content)
#             for chunk in chunks:
#                 all_chunks.append(Document(source=doc.source, page_content=chunk, metadata=doc.metadata))
#
#     if 'pdf' in file_types and documents['pdf'] is not None:
#         for doc in documents['pdf']:
#             chunks = split_by_braces_and_comma(doc.page_content)
#             for chunk in chunks:
#                 all_chunks.append(Document(source=doc.source, page_content=chunk, metadata=doc.metadata))
#
#     if 'csv' in file_types and documents['csv'] is not None:
#         for doc in documents['csv']:
#             chunks = split_by_braces_and_comma(doc.page_content)
#             for chunk in chunks:
#                 all_chunks.append(Document(source=doc.source, page_content=chunk, metadata=doc.metadata))
#
#     return all_chunks


chunks_res = split_docs_to_chunks(DOCS, ['txt', 'json', 'docx', 'csv'])

from tqdm import tqdm
import os


def get_chroma_vectorstore(documents, embeddings, persist_directory):
    if os.path.isdir(persist_directory) and os.listdir(persist_directory):
        print("Loading existing Chroma vectorstore...")
        vectorstore = Chroma(
            embedding_function=embeddings, persist_directory=persist_directory
        )

        existing_files = get_uploaded_filenames('local')
        uniq_sources_to_add = set(
            doc.metadata['source'] for doc in documents
            if doc.metadata['source'] not in existing_files
        )

        if uniq_sources_to_add:
            # Отслеживаем прогресс добавления новых документов
            print("Adding new documents to vectorstore...")
            docs_to_add = [doc for doc in documents if doc.metadata['source'] in uniq_sources_to_add]
            for doc in tqdm(docs_to_add, desc="Processing Documents", unit="doc"):
                vectorstore.add_documents(documents=[doc], embedding=embeddings)

            # Обновляем метаданные
            for filename in uniq_sources_to_add:
                add_filename_to_metadata('local', filename)
        else:
            print('No new documents found, skipping addition step.')

    else:
        print("Creating and indexing new Chroma vectorstore...")
        # Отслеживаем прогресс создания нового векторного хранилища
        for doc in tqdm(documents, desc="Indexing Documents", unit="doc"):
            vectorstore = Chroma.from_documents(
                documents=[doc],
                embedding=embeddings, persist_directory=persist_directory
            )

        # Сохраняем уникальные источники
        uniq_sources_to_add = set(doc.metadata['source'] for doc in documents)
        for filename in uniq_sources_to_add:
            add_filename_to_metadata('local', filename)

    return vectorstore



vectorstore = get_chroma_vectorstore(documents=chunks_res, embeddings=embeddings, persist_directory=CHROMA_PATH)
retriever = vectorstore.as_retriever(search_kwargs={"k": 3}, search_type='similarity')

#---------------------------------------------------------------------


def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


